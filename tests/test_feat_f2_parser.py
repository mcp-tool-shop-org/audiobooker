"""Tests for v2.1 INPUT features (FT-PARSE-001/003/004/005/007/008).

Covers:
- FT-PARSE-004 DOCX parsing (parse_docx) with python-docx mocked when absent.
- FT-PARSE-007 Markdown inline cleaning (strip_markdown_inline + parse_text gate).
- FT-PARSE-005 folder input (read_folder_chapters).
- FT-PARSE-003 TOC-driven EPUB splitting (parse_epub use_toc).
- FT-PARSE-008 Italian / Portuguese language profiles.
- FT-PARSE-001 pronunciation lexicon load/save (CSV + JSON, phoneme type).
"""

from __future__ import annotations

import sys
import types
from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# Availability helpers
# ---------------------------------------------------------------------------

def _has_ebooklib() -> bool:
    try:
        import ebooklib  # noqa: F401
        return True
    except ImportError:
        return False


def _has_python_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


requires_ebooklib = pytest.mark.skipif(
    not _has_ebooklib(),
    reason="ebooklib not installed",
)


# ===========================================================================
# FT-PARSE-004 — DOCX parsing
# ===========================================================================

class _FakeStyle:
    def __init__(self, name):
        self.name = name


class _FakePara:
    def __init__(self, text, style_name="Normal"):
        self.text = text
        self.style = _FakeStyle(style_name)


class _FakeCoreProps:
    def __init__(self, title="", author=""):
        self.title = title
        self.author = author


class _FakeDocument:
    def __init__(self, paragraphs, title="", author=""):
        self.paragraphs = paragraphs
        self.core_properties = _FakeCoreProps(title, author)


def _install_fake_docx(monkeypatch, document):
    """Install a fake ``docx`` module whose Document() returns ``document``."""
    fake = types.ModuleType("docx")

    def _Document(path):  # noqa: N802 — mirror python-docx API
        return document

    fake.Document = _Document
    monkeypatch.setitem(sys.modules, "docx", fake)
    return fake


def _make_docx_file(tmp_path: Path) -> Path:
    """Create a placeholder .docx file on disk (content is mocked, not read)."""
    p = tmp_path / "book.docx"
    p.write_bytes(b"PK\x03\x04placeholder")  # non-empty; never actually parsed
    return p


def _body(words: int) -> str:
    return " ".join(["word"] * words)


class TestParseDocx:
    """parse_docx — heading-style split, pattern fallback, metadata, shape."""

    def test_missing_python_docx_raises_with_install_hint(self, tmp_path, monkeypatch):
        """When python-docx is absent, ImportError names the pip install."""
        monkeypatch.setitem(sys.modules, "docx", None)  # force ImportError
        from audiobooker.parser.docx import parse_docx
        docx_file = _make_docx_file(tmp_path)
        with pytest.raises(ImportError) as exc:
            parse_docx(docx_file)
        assert "pip install python-docx" in str(exc.value)

    def test_splits_on_heading_styles(self, tmp_path, monkeypatch):
        """Heading 1 / Heading 2 / Title paragraphs start new chapters."""
        from audiobooker.parser.docx import parse_docx
        doc = _FakeDocument(
            paragraphs=[
                _FakePara("The Great Book", "Title"),
                _FakePara(_body(60), "Normal"),
                _FakePara("Chapter One", "Heading 1"),
                _FakePara(_body(60), "Normal"),
                _FakePara("A Subsection", "Heading 2"),
                _FakePara(_body(60), "Normal"),
            ],
            title="Core Title",
            author="A. Writer",
        )
        _install_fake_docx(monkeypatch, doc)
        meta, chapters = parse_docx(_make_docx_file(tmp_path), min_chapter_words=10)
        assert meta["title"] == "Core Title"
        assert meta["author"] == "A. Writer"
        assert [c.title for c in chapters] == ["The Great Book", "Chapter One", "A Subsection"]
        assert all(c.raw_text for c in chapters)

    def test_metadata_falls_back_to_filename(self, tmp_path, monkeypatch):
        """No core-property title -> metadata title is the file stem."""
        from audiobooker.parser.docx import parse_docx
        doc = _FakeDocument(
            paragraphs=[
                _FakePara("Heading", "Heading 1"),
                _FakePara(_body(60), "Normal"),
            ],
        )
        _install_fake_docx(monkeypatch, doc)
        meta, chapters = parse_docx(_make_docx_file(tmp_path), min_chapter_words=10)
        assert meta["title"] == "book"  # book.docx stem

    def test_pattern_fallback_when_no_heading_styles(self, tmp_path, monkeypatch):
        """With no styled headings, the profile's chapter_patterns split text."""
        from audiobooker.parser.docx import parse_docx
        from audiobooker.language.profile import get_profile
        doc = _FakeDocument(
            paragraphs=[
                _FakePara("Chapter 1", "Normal"),
                _FakePara(_body(60), "Normal"),
                _FakePara("Chapter 2", "Normal"),
                _FakePara(_body(60), "Normal"),
            ],
        )
        _install_fake_docx(monkeypatch, doc)
        meta, chapters = parse_docx(
            _make_docx_file(tmp_path),
            min_chapter_words=10,
            profile=get_profile("en"),
        )
        assert len(chapters) == 2

    def test_return_shape_matches_parse_epub(self, tmp_path, monkeypatch):
        """parse_docx returns (dict, list[Chapter]) like parse_epub."""
        from audiobooker.parser.docx import parse_docx
        from audiobooker.models import Chapter
        doc = _FakeDocument(
            paragraphs=[
                _FakePara("H", "Heading 1"),
                _FakePara(_body(60), "Normal"),
            ],
        )
        _install_fake_docx(monkeypatch, doc)
        result = parse_docx(_make_docx_file(tmp_path), min_chapter_words=10)
        assert isinstance(result, tuple) and len(result) == 2
        meta, chapters = result
        assert isinstance(meta, dict)
        assert isinstance(chapters, list)
        assert all(isinstance(c, Chapter) for c in chapters)

    def test_empty_document_raises(self, tmp_path, monkeypatch):
        """A document with no qualifying content raises ValueError."""
        from audiobooker.parser.docx import parse_docx
        doc = _FakeDocument(paragraphs=[_FakePara("", "Normal")])
        _install_fake_docx(monkeypatch, doc)
        with pytest.raises(ValueError):
            parse_docx(_make_docx_file(tmp_path), min_chapter_words=10)

    def test_missing_file_raises(self, tmp_path, monkeypatch):
        from audiobooker.parser.docx import parse_docx
        _install_fake_docx(monkeypatch, _FakeDocument(paragraphs=[]))
        with pytest.raises(FileNotFoundError):
            parse_docx(tmp_path / "nope.docx")


# ===========================================================================
# FT-PARSE-007 — Markdown inline cleaning
# ===========================================================================

class TestStripMarkdownInline:
    """strip_markdown_inline unwraps formatting to spoken text."""

    def test_unwraps_emphasis(self):
        from audiobooker.parser.text_cleaners import strip_markdown_inline
        out = strip_markdown_inline("This is **bold**, *em*, _em2_ and ~~gone~~.")
        assert out == "This is bold, em, em2 and gone."

    def test_links_and_images(self):
        from audiobooker.parser.text_cleaners import strip_markdown_inline
        out = strip_markdown_inline("See [the site](http://x.com) and ![a cat](cat.png).")
        assert "the site" in out
        assert "a cat" in out
        assert "http://x.com" not in out
        assert "cat.png" not in out

    def test_inline_and_fenced_code(self):
        from audiobooker.parser.text_cleaners import strip_markdown_inline
        md = "Run `make build` now.\n\n```python\nprint('x')\n```\n\nDone."
        out = strip_markdown_inline(md)
        assert "make build" in out
        assert "`" not in out
        assert "print('x')" not in out
        assert "Done." in out

    def test_list_markers_and_blockquotes(self):
        from audiobooker.parser.text_cleaners import strip_markdown_inline
        md = "- first\n* second\n1. third\n> a quote"
        out = strip_markdown_inline(md)
        lines = [ln for ln in out.splitlines() if ln.strip()]
        assert lines == ["first", "second", "third", "a quote"]

    def test_plain_text_unchanged_by_cleaner(self):
        """A reference-form link still unwraps, but no markup means no change."""
        from audiobooker.parser.text_cleaners import strip_markdown_inline
        assert strip_markdown_inline("Just a plain sentence.") == "Just a plain sentence."


class TestParseTextMarkdownGate:
    """parse_text applies markdown stripping ONLY for .md/.markdown."""

    def test_md_file_is_stripped(self, tmp_path):
        from audiobooker.parser.text import parse_text
        p = tmp_path / "story.md"
        p.write_text(
            "This **bold** word and a [link](http://x.com) plus more "
            "words to clear the threshold and form a real paragraph here.",
            encoding="utf-8",
        )
        _meta, chapters = parse_text(p)
        text = chapters[0].raw_text
        assert "**" not in text
        assert "http://x.com" not in text
        assert "bold" in text

    def test_txt_file_keeps_asterisks(self, tmp_path):
        """Plain .txt with legitimate asterisks is NOT markdown-stripped."""
        from audiobooker.parser.text import parse_text
        p = tmp_path / "story.txt"
        p.write_text(
            "The hero said the *real* secret out loud, and then *another* "
            "thing happened to fill out the paragraph with enough words.",
            encoding="utf-8",
        )
        _meta, chapters = parse_text(p)
        text = chapters[0].raw_text
        # Asterisks survive in plain text.
        assert "*real*" in text
        assert "*another*" in text

    def test_markdown_extension_variant(self, tmp_path):
        from audiobooker.parser.text import parse_text
        p = tmp_path / "story.markdown"
        p.write_text(
            "A `code` token and **strong** words should be unwrapped while "
            "this paragraph stays nicely above the minimum word threshold.",
            encoding="utf-8",
        )
        _meta, chapters = parse_text(p)
        text = chapters[0].raw_text
        assert "`" not in text
        assert "**" not in text


# ===========================================================================
# FT-PARSE-005 — Folder input
# ===========================================================================

class TestReadFolderChapters:
    """read_folder_chapters orders, titles, and reads per-chapter files."""

    def test_natural_sort_of_numeric_prefixes(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "1_intro.txt").write_text("alpha", encoding="utf-8")
        (tmp_path / "2_middle.txt").write_text("beta", encoding="utf-8")
        (tmp_path / "10_end.txt").write_text("gamma", encoding="utf-8")
        chapters = read_folder_chapters(tmp_path)
        # 2 must sort before 10 (numeric, not lexicographic).
        assert [t for t, _ in chapters] == ["Intro", "Middle", "End"]
        assert [c for _, c in chapters] == ["alpha", "beta", "gamma"]

    def test_dotted_and_bare_numeric_prefixes(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "01.txt").write_text("one", encoding="utf-8")
        (tmp_path / "2. The Road.txt").write_text("two", encoding="utf-8")
        chapters = read_folder_chapters(tmp_path)
        assert chapters[0][0] == "Chapter 1"      # pure-number stem
        assert chapters[1][0] == "The Road"       # "2. The Road" -> "The Road"

    def test_frontmatter_title_wins(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "01_chapter.md").write_text(
            "---\ntitle: A Custom Title\n---\nThe body of the chapter.",
            encoding="utf-8",
        )
        chapters = read_folder_chapters(tmp_path)
        assert chapters[0][0] == "A Custom Title"
        assert "The body of the chapter." in chapters[0][1]

    def test_markdown_files_are_stripped(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "01_ch.md").write_text("This is **bold** content.", encoding="utf-8")
        chapters = read_folder_chapters(tmp_path)
        assert "**" not in chapters[0][1]
        assert "bold" in chapters[0][1]

    def test_pattern_filters_extensions(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "01.txt").write_text("text file", encoding="utf-8")
        (tmp_path / "02.md").write_text("md file", encoding="utf-8")
        (tmp_path / "notes.log").write_text("ignored", encoding="utf-8")
        only_txt = read_folder_chapters(tmp_path, pattern="*.txt")
        assert len(only_txt) == 1
        both = read_folder_chapters(tmp_path, pattern="*.txt;*.md")
        assert len(both) == 2

    def test_missing_folder_raises(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        with pytest.raises(FileNotFoundError):
            read_folder_chapters(tmp_path / "nope")

    def test_no_matching_files_raises(self, tmp_path):
        from audiobooker.parser.text import read_folder_chapters
        (tmp_path / "readme.rst").write_text("x", encoding="utf-8")
        with pytest.raises(ValueError):
            read_folder_chapters(tmp_path)


# ===========================================================================
# FT-PARSE-003 — TOC-driven EPUB splitting
# ===========================================================================

def _make_toc_epub(tmp_path: Path) -> Path:
    """Build an EPUB with ONE document holding three anchored sections and a
    TOC of Links pointing at those anchors."""
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("toc-book-001")
    book.set_title("TOC Book")
    book.set_language("en")
    book.add_author("Anchor Author")

    body = (
        "<html><body>"
        '<h1 id="c1">First Movement</h1>'
        "<p>" + _body(80) + "</p>"
        '<h1 id="c2">Second Movement</h1>'
        "<p>" + _body(80) + "</p>"
        '<h1 id="c3">Third Movement</h1>'
        "<p>" + _body(80) + "</p>"
        "</body></html>"
    )
    doc = epub.EpubHtml(title="Whole", file_name="content.xhtml", lang="en")
    doc.content = body
    book.add_item(doc)

    book.toc = (
        epub.Link("content.xhtml#c1", "First Movement", "c1"),
        epub.Link("content.xhtml#c2", "Second Movement", "c2"),
        epub.Link("content.xhtml#c3", "Third Movement", "c3"),
    )
    book.spine = ["nav", doc]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    path = tmp_path / "toc.epub"
    epub.write_epub(str(path), book)
    return path


def _make_multi_doc_epub(tmp_path: Path) -> Path:
    """Build an EPUB with multiple documents and NO usable TOC, to exercise the
    spine fallback path."""
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("spine-book-001")
    book.set_title("Spine Book")
    book.set_language("en")
    book.add_author("Spine Author")

    docs = []
    for i in range(3):
        ch = epub.EpubHtml(title=f"Doc {i}", file_name=f"d{i}.xhtml", lang="en")
        ch.content = f"<html><body><h1>Doc {i}</h1><p>{_body(80)}</p></body></html>"
        book.add_item(ch)
        docs.append(ch)

    book.toc = ()  # no TOC at all
    book.spine = ["nav"] + docs
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    path = tmp_path / "spine.epub"
    epub.write_epub(str(path), book)
    return path


@requires_ebooklib
class TestTocDrivenSplit:
    """parse_epub use_toc — TOC slicing, spine fallback, and 'off' force."""

    def test_toc_split_produces_titled_chapters(self, tmp_path, caplog):
        import logging
        from audiobooker.parser.epub import parse_epub
        epub_path = _make_toc_epub(tmp_path)
        with caplog.at_level(logging.INFO, logger="audiobooker.parser"):
            _meta, chapters = parse_epub(epub_path, min_chapter_words=10, use_toc="auto")
        titles = [c.title for c in chapters]
        assert titles == ["First Movement", "Second Movement", "Third Movement"]
        # Anchor slicing keeps each movement's text separate.
        assert "Second Movement" not in chapters[0].raw_text
        assert "Third Movement" not in chapters[1].raw_text
        # Observability reports the TOC path.
        assert any("pattern=toc" in r.getMessage() for r in caplog.records)

    def test_use_toc_off_forces_spine(self, tmp_path, caplog):
        import logging
        from audiobooker.parser.epub import parse_epub
        epub_path = _make_toc_epub(tmp_path)
        with caplog.at_level(logging.INFO, logger="audiobooker.parser"):
            _meta, chapters = parse_epub(epub_path, min_chapter_words=10, use_toc="off")
        # TOC ignored: the single content document is NOT anchor-sliced, so the
        # three movements collapse into one chapter (plus possibly a nav doc).
        # The key signal is that spine splitting ran and TOC titles did not.
        titles = [c.title for c in chapters]
        assert titles != ["First Movement", "Second Movement", "Third Movement"]
        assert any("pattern=spine/structure" in r.getMessage() for r in caplog.records)

    def test_no_toc_falls_back_to_spine(self, tmp_path):
        from audiobooker.parser.epub import parse_epub
        epub_path = _make_multi_doc_epub(tmp_path)
        _meta, chapters = parse_epub(epub_path, min_chapter_words=10, use_toc="auto")
        # The three real content documents each become a chapter (the EPUB's
        # generated nav document may add one more, so assert the content set is
        # present rather than an exact count).
        titles = [c.title for c in chapters]
        for expected in ("Doc 0", "Doc 1", "Doc 2"):
            assert expected in titles

    def test_invalid_use_toc_raises(self, tmp_path):
        from audiobooker.parser.epub import parse_epub
        epub_path = _make_multi_doc_epub(tmp_path)
        with pytest.raises(ValueError):
            parse_epub(epub_path, use_toc="bogus")

    def test_default_use_toc_is_auto(self, tmp_path):
        """Default (no use_toc arg) splits the anchored single-doc book by TOC."""
        from audiobooker.parser.epub import parse_epub
        epub_path = _make_toc_epub(tmp_path)
        _meta, chapters = parse_epub(epub_path, min_chapter_words=10)
        assert len(chapters) == 3


# ===========================================================================
# FT-PARSE-008 — Italian / Portuguese profiles
# ===========================================================================

class TestItalianProfile:
    def test_registered_and_basic_fields(self):
        from audiobooker.language.profile import get_profile, available_profiles
        assert "it" in available_profiles()
        it = get_profile("it")
        assert it.name == "Italian"
        assert "disse" in it.speaker_verbs
        assert "chiese" in it.speaker_verbs
        assert "rispose" in it.speaker_verbs

    def test_dialogue_quotes_include_guillemets(self):
        from audiobooker.language.profile import get_profile
        it = get_profile("it")
        pairs = set(it.dialogue_quotes) | set(it.smart_quotes)
        assert ("«", "»") in pairs           # «»
        assert ("“", "”") in pairs           # “”

    def test_accented_name_and_chapter_pattern(self):
        import re
        from audiobooker.language.profile import get_profile
        it = get_profile("it")
        assert it.is_valid_name("Niccolò")
        assert it.is_valid_name("Beatrice")
        assert re.match(it.chapter_patterns[0], "Capitolo 1: L'inizio")
        assert re.match(it.chapter_patterns[1], "Parte 2")

    def test_emotion_hints(self):
        from audiobooker.language.profile import get_profile
        it = get_profile("it")
        assert it.emotion_hints.get("sussurrò") == "whisper"
        assert it.emotion_hints.get("gridò") == "angry"


class TestPortugueseProfile:
    def test_registered_and_basic_fields(self):
        from audiobooker.language.profile import get_profile, available_profiles
        assert "pt" in available_profiles()
        pt = get_profile("pt")
        assert pt.name == "Portuguese"
        assert "disse" in pt.speaker_verbs
        assert "perguntou" in pt.speaker_verbs
        assert "respondeu" in pt.speaker_verbs

    def test_dialogue_uses_em_dash_and_guillemets(self):
        from audiobooker.language.profile import get_profile
        pt = get_profile("pt")
        openers = {pair[0] for pair in pt.dialogue_quotes}
        assert "—" in openers                      # em-dash travessão
        assert "«" in openers                      # «

    def test_accented_name_and_chapter_pattern(self):
        import re
        from audiobooker.language.profile import get_profile
        pt = get_profile("pt")
        assert pt.is_valid_name("João")
        assert pt.is_valid_name("Conceição")
        assert re.match(pt.chapter_patterns[0], "Capítulo 1")
        assert re.match(pt.chapter_patterns[1], "Parte 3")

    def test_emotion_hints(self):
        from audiobooker.language.profile import get_profile
        pt = get_profile("pt")
        assert pt.emotion_hints.get("sussurrou") == "whisper"
        assert pt.emotion_hints.get("gritou") == "angry"


# ===========================================================================
# FT-PARSE-001 — Pronunciation lexicon load/save
# ===========================================================================

class TestLexicon:
    def test_csv_load(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon
        p = tmp_path / "lex.csv"
        p.write_text(
            "word,replacement,type\nGandalf,GAHN-dalf,text\nSmaug,smOWg,phoneme\n",
            encoding="utf-8",
        )
        lex = load_lexicon(p)
        assert lex["Gandalf"] == "GAHN-dalf"
        assert lex["Smaug"] == "smOWg"
        assert lex.types["Smaug"] == "phoneme"

    def test_csv_without_header(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon
        p = tmp_path / "lex.csv"
        p.write_text("Frodo,FROH-doh\nArwen,AR-wen\n", encoding="utf-8")
        lex = load_lexicon(p)
        assert lex["Frodo"] == "FROH-doh"
        assert lex["Arwen"] == "AR-wen"

    def test_json_flat_load(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon
        p = tmp_path / "lex.json"
        p.write_text('{"Bilbo": "BILL-boh"}', encoding="utf-8")
        lex = load_lexicon(p)
        assert lex["Bilbo"] == "BILL-boh"
        assert lex.types["Bilbo"] == "text"

    def test_json_typed_load(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon
        p = tmp_path / "lex.json"
        p.write_text(
            '{"Smaug": {"replacement": "smOWg", "type": "phoneme"}}',
            encoding="utf-8",
        )
        lex = load_lexicon(p)
        assert lex["Smaug"] == "smOWg"
        assert lex.types["Smaug"] == "phoneme"

    def test_phoneme_type_round_trips_csv_to_json(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon, save_lexicon
        src = tmp_path / "src.csv"
        src.write_text(
            "word,replacement,type\nA,aa,text\nB,bb,phoneme\n", encoding="utf-8"
        )
        lex = load_lexicon(src)
        dst = tmp_path / "dst.json"
        save_lexicon(dst, lex)
        reloaded = load_lexicon(dst)
        assert reloaded.types["B"] == "phoneme"
        assert reloaded.types["A"] == "text"

    def test_save_csv_has_no_blank_lines(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon, save_lexicon
        src = tmp_path / "src.json"
        src.write_text('{"X": "xx", "Y": "yy"}', encoding="utf-8")
        lex = load_lexicon(src)
        dst = tmp_path / "dst.csv"
        save_lexicon(dst, lex)
        lines = dst.read_text(encoding="utf-8").splitlines()
        assert lines == ["word,replacement,type", "X,xx,text", "Y,yy,text"]

    def test_lexicon_usable_as_overrides(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon, apply_pronunciation_overrides
        p = tmp_path / "lex.csv"
        p.write_text("word,replacement\nGandalf,GAHN-dalf\n", encoding="utf-8")
        lex = load_lexicon(p)
        out = apply_pronunciation_overrides("Then Gandalf spoke.", lex)
        assert out == "Then GAHN-dalf spoke."

    def test_missing_lexicon_file_raises(self, tmp_path):
        from audiobooker.parser.text_cleaners import load_lexicon
        with pytest.raises(FileNotFoundError):
            load_lexicon(tmp_path / "nope.csv")
